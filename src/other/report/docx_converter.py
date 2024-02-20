import docx
from docx.shared import Pt, RGBColor


class DocxConverter:
    def __init__(self, data: dict, path: str):
        self._data = data
        self._path = path
        self._document = None
        self._default_styles = {
            'Normal': {'font': 'Times New Roman', 'size': 14, 'color': (0, 0, 0)},
            'Title': {'font': 'Times New Roman', 'size': 22, 'color': (0, 0, 0)},
            'Heading 1': {'font': 'Times New Roman', 'size': 18, 'bold': True, 'color': (0, 0, 0)},
            'Heading 2': {'font': 'Times New Roman', 'size': 16, 'bold': True, 'color': (0, 0, 0)},
            'Body Text': {'font': 'Consolas', 'size': 12, 'color': (0, 0, 0)},
        }

    def open_file(self):
        self._document = docx.Document()

    def close_file(self):
        self._document.save(self._path)

    def _convert_main_document(self, data: dict):
        if data.get('name'):
            self._document.add_heading(data.get('name'), level=0)
        for el in data.get('children', []):
            self._select_converter(el)

    def _convert_section(self, data: dict):
        self._document.add_heading(data.get('name'), level=1)
        for el in data.get('children', []):
            self._select_converter(el)

    def _convert_subsection(self, data: dict):
        self._document.add_heading(data.get('name'), level=2)
        for el in data.get('children', []):
            self._select_converter(el)

    def _convert_text_area(self, data: dict):
        self._document.add_paragraph(data.get('text'))

    def _convert_code_area(self, data: dict):
        self._document.add_paragraph(data.get('text'), 'Body Text')

    def _convert_list(self, data: dict):
        style = 'List Bullet' if data.get('type') == 0 else 'List Number'
        for el in data.get('children', []):
            self._document.add_paragraph(el.get('text'), style=style)

    def _select_converter(self, data: dict):
        if data.get('__class__') == 'ReportMainDocument':
            self._convert_main_document(data)
        if data.get('__class__') == 'ReportSection':
            self._convert_section(data)
        if data.get('__class__') == 'ReportSubSection':
            self._convert_subsection(data)
        if data.get('__class__') == 'ReportTextEdit':
            self._convert_text_area(data)
        if data.get('__class__') == 'ReportCodeEdit':
            self._convert_code_area(data)
        if data.get('__class__') == 'ReportListWidget':
            self._convert_list(data)

    def _set_styles(self, data: dict):
        for key, item in data.items():
            style = self._document.styles[key]
            if 'font' in item:
                style.font.name = item['font']
            if 'size' in item:
                style.font.size = Pt(item['size'])
            if 'color' in item:
                style.font.color.rgb = RGBColor(*item.get('color'))
            style.font.italic = item.get('italic', False)
            style.font.bold = item.get('bold', False)
            style.font.underline = item.get('underline', False)
            style.font.strike = item.get('strike', False)

    def convert(self):
        self.open_file()
        if 'styles' in self._data:
            self._set_styles(self._data['styles'])
        else:
            self._set_styles(self._default_styles)
        self._convert_main_document(self._data)
        self.close_file()
