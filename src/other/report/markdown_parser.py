import os
import re
import sys
from time import sleep

import docx
from PIL import Image
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Pt, RGBColor, Mm
from docx.oxml.ns import qn
from htmldocx import HtmlToDocx
from lxml import etree
from requests import post
import docx2pdf
from cairosvg import svg2png
import latex2mathml.converter as latex_converter
from pypdf import PdfMerger, PdfReader

from src import config
from src.backend.commands import read_file
from src.other.report.docx_api import docx_to_pdf_by_api

DEFAULT_STILES = {
    'Normal': {'font': 'Times New Roman', 'size': 14, 'color': (0, 0, 0)},
    'Title': {'font': 'Times New Roman', 'size': 22, 'color': (0, 0, 0)},
    'Heading 1': {'font': 'Times New Roman', 'size': 16, 'bold': True, 'color': (0, 0, 0)},
    'Heading 2': {'font': 'Times New Roman', 'size': 15, 'bold': True, 'color': (0, 0, 0)},
    'Heading 3': {'font': 'Times New Roman', 'size': 14, 'bold': True, 'underline': True, 'color': (0, 0, 0)},
    'Heading 4': {'font': 'Times New Roman', 'size': 14, 'bold': True, 'color': (0, 0, 0)},
    'Heading 5': {'font': 'Times New Roman', 'size': 14, 'italic': True, 'underline': True, 'color': (0, 0, 0)},
    'Heading 6': {'font': 'Times New Roman', 'size': 14, 'italic': True, 'color': (0, 0, 0)},
    'Body Text': {'font': 'Consolas', 'size': 11, 'color': (0, 0, 0)},
}

DEFAULT_PROPERTIES = {
    'margins': (30, 20, 10, 20),
    'first_line_indent': 12.5,
}

_YES = 1
_NO = 0
_UNKNOWN = 2


class MarkdownParser:
    def __init__(self, bm, text: str, dist: str, to_pdf="", styles=None, properties=None):
        self.bm = bm
        self.text = text
        self.dist = dist
        self.to_pdf = to_pdf

        self.document = docx.Document()

        self.html_parser = HtmlToDocx()
        self._pdf_merger = PdfMerger()

        self.styles = dict()
        self._set_styles(styles if styles else DEFAULT_STILES)
        self.properties = properties if properties else DEFAULT_PROPERTIES

        self._lines = self.text.split('\n')
        self._current_line = -1

        self._list_id = 9

        self._pdf_to_merge = []
        self._last_paragraph = None
        self._italic = _NO
        self._bold = _NO
        self._code = _NO

    def _next_line(self):
        self._current_line += 1
        if self._current_line >= len(self._lines):
            return None
        return self._lines[self._current_line]

    def return_line(self):
        self._current_line -= 1

    def _have_lines(self):
        return self._current_line < len(self._lines) - 1

    def convert(self):
        while (line := self._next_line()) is not None:
            if self.parse_image(line):
                continue
            if self.parse_code(line):
                continue
            if self.parse_toc(line):
                continue
            if self.parse_author(line):
                continue
            if self.parse_page_size(line):
                continue
            if self.parse_pdf_including(line):
                continue
            if self.parse_header(line):
                continue
            if self.parse_bullet_list(line):
                continue
            if self.parse_num_list(line):
                continue
            if self.parse_page_break(line):
                continue
            if self.parse_line_break(line):
                continue
            if self.parse_formula(line):
                continue
            if self.parse_simple_formula(line):
                continue
            if self.parse_tests(line):
                continue
            if self.parse_table(line):
                continue
            if self.parse_paragraph(line):
                continue

        self.set_margins()
        self._add_page_numbers()
        self.document.save(self.dist)
        if self.to_pdf:
            self.convert_to_pdf()

    def parse_paragraph(self, line):
        if not line.strip():
            return False
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        margin = count_in_start(line, ' ') // 4
        if margin:
            paragraph.paragraph_format.left_indent = Mm(12.5 * margin)
        else:
            paragraph.paragraph_format.first_line_indent = Mm(self.properties.get('first_line_indent', 0))
        self.add_runs(paragraph, line.strip())
        while (line := self._next_line()) is not None and line.strip():
            if self.parse_simple_formula(line, paragraph):
                continue
            if line.lstrip().startswith('- ') or line.lstrip().startswith('[') or line.lstrip(' 1234567890').startswith(
                    '. ') or line.lstrip().startswith('#'):
                self.return_line()
                break
            self.add_runs(paragraph, line.strip())
        return True

    def parse_header(self, line):
        if not line.startswith('#'):
            return False
        p = self.document.add_heading(line.lstrip('#').strip(), count_in_start(line, '#'))
        p.paragraph_format.first_line_indent = Mm(self.properties.get('first_line_indent', 0))
        return True

    def parse_bullet_list(self, line):
        if not line.strip().startswith('- '):
            return False
        paragraph = self.document.add_paragraph(style="List Bullet")

        margin = count_in_start(line, ' ')
        if margin:
            paragraph.paragraph_format.left_indent = Mm(12.5 * (margin // 4))

        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.add_runs(paragraph, line.lstrip().lstrip('-').strip())
        while (line := self._next_line()) is not None and line.strip():
            if line.strip().startswith('- '):
                level = (count_in_start(line, ' ') - margin) // 2
                paragraph = self.document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                if margin:
                    paragraph.paragraph_format.left_indent = Mm(12.5 * (margin // 4))
                MarkdownParser._billet_list(paragraph, level)
            self.add_runs(paragraph, line.lstrip().lstrip('-').strip())
        return True

    def parse_num_list(self, line):
        if not line.strip().lstrip('1234567890').startswith('. ') or line.strip().startswith('.'):
            return False
        level = count_in_start(line, ' ') // 3
        paragraph = self.document.add_paragraph(style="List Number")
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self._num_list(paragraph, level, new_list=line.strip().startswith('1.'))
        self.add_runs(paragraph, line.strip().lstrip('1234567890.').lstrip())
        while (line := self._next_line()) is not None and line.strip():
            if line.strip().lstrip('1234567890').startswith('. ') and not line.strip().startswith('.'):
                level = count_in_start(line, ' ') // 3
                paragraph = self.document.add_paragraph(style="List Number")
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                self._num_list(paragraph, level)
            self.add_runs(paragraph, line.strip().lstrip('1234567890.').lstrip())
        return True

    def parse_page_break(self, line):
        if not line.startswith('[page-break]: <>'):
            return False
        self.document.add_page_break()
        return True

    def parse_line_break(self, line):
        if not line.startswith('[line-break]: <>') and not line.startswith('[br]: <>'):
            return False
        self.document.add_paragraph()
        return True

    def parse_formula(self, line, paragraph=None):
        if line.startswith('[formula-start]: <>'):
            lines = []
            while (line := self._next_line()) is not None and not line.startswith('[formula-end]: <>'):
                lines.append(line)
            self.convert_formula('\n'.join(lines), paragraph)
            return True
        return False

    def parse_simple_formula(self, line, paragraph=None):
        if line.startswith('[formula]: <> ('):
            self.convert_formula(line[len('[formula]: <> ('):-1], paragraph)
            return True
        return False

    def parse_toc(self, line):
        if line.startswith('[table-of-content]: <>'):
            self.add_table_of_content()
            return True
        return False

    def parse_author(self, line):
        if line.startswith(tag := '[author]: <> ('):
            self.document.core_properties.author = line.strip()[len(tag):-1]
            return True
        return False

    def parse_pdf_including(self, line):
        if line.startswith(tag := '[include-pdf]: <> ('):
            self._pdf_to_merge.append(line.strip()[len(tag):-1])
            return True
        return False

    def parse_page_size(self, line):
        if line.startswith(tag := '[page-size]: <> ('):
            width, height = line.strip()[len(tag):-1].split()
            for section in self.document.sections:
                section.page_width = Mm(float(width.strip()))
                section.page_height = Mm(float(height.strip()))
            return True
        return False

    def parse_image(self, line):
        if not re.match(r"!\[[\w.=\\/:]*]\([\w.\\/:]+\)", line.strip()):
            return False
        default_text, image_path = line.strip()[2:line.index(')')].split('](')
        if image_path.endswith('.svg'):
            svg2png(url=image_path, write_to=(image_path := f"{self.bm._sm.temp_dir()}/image.png"))
        img = Image.open(image_path)
        height, width = img.height, img.width
        if default_text.startswith('height='):
            h = int(default_text.lstrip('height='))
            width = width * h // height
            height = h
        elif default_text.startswith('width='):
            w = int(default_text.lstrip('width='))
            height = height * w // width
            width = w
        elif width > 170:
            height = height * 170 // width
            width = 170
        img.close()
        try:
            self.document.add_picture(image_path, width=Mm(width), height=Mm(height))
        except Exception:
            self.document.add_paragraph(default_text)
        return True

    def parse_code(self, line):
        if not line.startswith('```'):
            return False
        lexer = line.lstrip('```').strip()
        code_lines = []
        while (line := self._next_line()) is not None and not line.endswith('```'):
            code_lines.append(line)
        self.highlight_code('\n'.join(code_lines), lexer)
        return True

    def parse_table(self, line: str):
        if not line.startswith('|'):
            return False
        table_columns = line.strip().count('|') - 1
        table_rows = 1
        table = self.document.add_table(rows=1, cols=table_columns, style='Table Grid')
        for i, text in enumerate(line.strip()[1:-1].split('|')):
            table.cell(0, i).text = text.strip()
        while (line := self._next_line()) is not None and line.startswith('|'):
            if line.strip().strip('-|'):
                table.add_row()
                for i, text in enumerate(line.strip()[1:-1].split('|')):
                    if not text.strip() and i > 1:
                        text = table.cell(table_rows, i - 1).text
                        table.cell(table_rows, i).merge(table.cell(table_rows, i - 1)).text = text
                    table.cell(table_rows, i).text = text.strip()
                table_rows += 1
        return True

    def add_runs(self, paragraph, line):
        if self.parse_simple_formula(line, paragraph):
            self._add_runs(paragraph, ' ')
            return
        self._add_runs(paragraph, line + ' ')

    def _add_runs(self, paragraph, line, bold=_UNKNOWN, italic=_UNKNOWN, code=_UNKNOWN):
        def split_line(line_to_split: str, sep: str):
            lst = line_to_split.split(sep)
            i = 0
            while i < len(lst) - 1:
                if (len(lst[i]) - len(lst[i].rstrip('\\'))) % 2:
                    lst[i] = lst[i + 1]
                    lst.pop(i + 1)
                i += 1
            return lst

        if bold != _UNKNOWN and italic != _UNKNOWN and code != _UNKNOWN:
            run = paragraph.add_run(line.replace('\\*', '*').replace('\\`', '`').replace('\\\\', '\\'))
            run.font.bold = bool(bold)
            run.font.italic = bool(italic)
            if code:
                run.font.name = 'Consolas'
            self._last_paragraph = paragraph
            self._bold, self._italic, self._code = bold, italic, code
        elif bold == _UNKNOWN:
            flag = self._bold if paragraph == self._last_paragraph else _NO
            for el in split_line(line, '**'):
                self._add_runs(paragraph, el, flag, italic, code)
                flag = _NO if flag == _YES else _YES
        elif italic == _UNKNOWN:
            flag = self._italic if paragraph == self._last_paragraph else _NO
            for el in split_line(line, '*'):
                self._add_runs(paragraph, el, bold, flag, code)
                flag = _NO if flag == _YES else _YES
        elif code == _UNKNOWN:
            flag = self._code if paragraph == self._last_paragraph else _NO
            for el in split_line(line, '`'):
                self._add_runs(paragraph, el, bold, italic, flag)
                flag = _NO if flag == _YES else _YES

    def add_runs_old(self, paragraph, line):
        words = line.split(' ')
        code = False
        bold = False
        italic = False
        run_code = []

        for word in words:
            text = word
            if run_code:
                if text.endswith(')$'):
                    text = text[:-1]
                    run_code.append(text)
                    text = str(eval(' '.join(run_code)))
                    run_code.clear()
                else:
                    run_code.append(text + ' ')
            else:
                if text == '*' or text == '**':
                    pass
                elif text.startswith('`'):
                    code = True
                    text = text[1:]
                elif text.startswith('***'):
                    bold = True
                    italic = True
                    text = text[3:]
                elif text.startswith('**'):
                    bold = True
                    text = text[2:]
                elif text.startswith('*'):
                    italic = True
                    text = text[1:]
                elif text.startswith('$('):
                    text = text[2:]
                    run_code.append(text)

                elif text == '*' or text == '**':
                    pass
                elif text.endswith('`') and code:
                    text = text[:-1]
                elif text.endswith('***') and bold and italic:
                    text = text[:-3]
                elif text.endswith('**') and bold:
                    text = text[:-2]
                elif text.endswith('*') and italic:
                    text = text[:-1]

            if not run_code:
                run = paragraph.add_run(text)
                if code:
                    run.font.name = 'Courier'
                if italic:
                    run.font.italic = True
                if bold:
                    run.font.bold = True
                paragraph.add_run(' ')

                if word.endswith('`'):
                    code = False
                elif word.endswith('***'):
                    bold = False
                    italic = False
                elif word.endswith('**'):
                    bold = False
                elif word.endswith('*'):
                    italic = False

    def parse_tests(self, line):
        if line.startswith('[tests]: <>'):
            flags = set(map(str.upper, line[len('[tests]: <>'):].strip().strip('()').split()))
            lst = ["Описание", "Входные данные"]
            if 'EXPECTED_OUTPUT' in flags:
                lst.append("Ожидаемый вывод")
            if 'REAL_OUTPUT' in flags:
                lst.append("Фактический вывод")
            if 'STATUS' in flags:
                lst.append("Результат")
            tests = self.bm._func_tests['pos'] + self.bm._func_tests['neg']
            table = self.document.add_table(rows=len(tests) + 1, cols=len(lst), style='Table Grid')
            if 'RUN' in flags:
                looper = self.bm.start_testing()
                while not looper.isFinished():
                    sleep(0.1)
            for i, el in enumerate(lst):
                table.cell(0, i).text = el
            for i, test in enumerate(self._prepare_tests_data(tests, flags)):
                for j, el in enumerate(test):
                    table.cell(i + 1, j).text = el
            return True
        return False

    @staticmethod
    def _prepare_tests_data(tests, flags):
        for test in tests:
            lst = [test.get('desc', '').strip(), test.get('in', '').strip()]
            if 'EXPECTED_OUTPUT' in flags:
                lst.append(test.get('out', '').strip())
            if 'REAL_OUTPUT' in flags:
                if 'RUN' in flags:
                    lst.append(test.prog_out.get('STDOUT', '').strip())
                else:
                    lst.append(test.get('out', ''))
            if 'STATUS' in flags:
                lst.append("OK" if 'RUN' not in flags or test.res() else "FAIL")
            yield lst

    def set_margins(self):
        sections = self.document.sections
        for section in sections:
            section.left_margin = Mm(self.properties['margins'][0])
            section.top_margin = Mm(self.properties['margins'][1])
            section.right_margin = Mm(self.properties['margins'][2])
            section.bottom_margin = Mm(self.properties['margins'][3])

    def _set_styles(self, data: dict):
        for key, item in data.items():
            if key in self.document.styles:
                style = self.document.styles[key]
            else:
                style = self.document.styles.add_style(key, style_type=data.get('type', 1))
            if 'font' in item:
                style.font.name = item['font']
            if 'size' in item:
                style.font.size = Pt(item['size'])
            if 'color' in item:
                style.font.color.rgb = RGBColor(*item.get('color'))
            style.font.italic = item.get('italic', False)
            style.font.bold = item.get('bold', False)
            style.font.underline = item.get('underline', False)
            style.font.strike = item.get('strike', False)

            rFonts = style.element.rPr.rFonts
            rFonts.set(qn("w:asciiTheme"), item['font'])
            rFonts.set(qn("w:hAnsiTheme"), item['font'])

            self.styles[key] = style

    def highlight_code(self, code: str, lexer: str = 'c'):
        def delete_paragraph(paragraph):
            p = paragraph._element
            p.getparent().remove(p)
            paragraph._p = paragraph._element = None

        req = {
            'code': code.strip(),
            'lexer': lexer,
            'options': [],
            'style': 'colorful',
            'linenos': True
        }
        try:
            if not lexer:
                raise Exception
            res = post('http://hilite.me/api', req)
            if res.status_code >= 400:
                raise Exception(str(res.status_code))
            text = res.text
        except Exception as ex:
            table = self.document.add_table(rows=1, cols=2)
            table.cell(0, 0).paragraphs[0].text = '\n'.join(map(str, range(1, code.strip().count('\n') + 2)))
            table.cell(0, 0).paragraphs[0].style = 'Body Text'
            table.cell(0, 1).paragraphs[0].text = code.strip()
            table.cell(0, 1).paragraphs[0].style = 'Body Text'
        else:
            self.html_parser.add_html_to_document(text, self.document)
            table = self.document.tables[-1]
            delete_paragraph(self.document.paragraphs[-1])
        table.style = 'Table Grid'
        for row in table.rows:
            row.cells[0].width = Mm(10)
            row.cells[1].width = Mm(167)
        table.cell(0, 1).paragraphs[0].runs[-1].text = table.cell(0, 1).paragraphs[0].runs[-1].text.rstrip()
        for run in table.cell(0, 0).paragraphs[0].runs:
            run.font.size = Pt(11)
            run.font.name = 'Consolas'
        for run in table.cell(0, 1).paragraphs[0].runs:
            run.font.size = Pt(11)
            run.font.name = 'Consolas'

    def _add_page_numbers(self):
        def create_element(name):
            return OxmlElement(name)

        def create_attribute(element, name, value):
            element.set(ns.qn(name), value)

        def add_page_number(paragraph):
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            page_num_run = paragraph.add_run()

            fldChar1 = create_element('w:fldChar')
            create_attribute(fldChar1, 'w:fldCharType', 'begin')

            instrText = create_element('w:instrText')
            create_attribute(instrText, 'xml:space', 'preserve')
            instrText.text = "PAGE"

            fldChar2 = create_element('w:fldChar')
            create_attribute(fldChar2, 'w:fldCharType', 'end')

            page_num_run._r.append(fldChar1)
            page_num_run._r.append(instrText)
            page_num_run._r.append(fldChar2)

        add_page_number(self.document.sections[0].footer.paragraphs[0])

    def _num_list(self, paragraph, level=0, new_list=False):
        if new_list:
            self._list_id += 1
        num_pr = OxmlElement('w:numPr')

        ilvl = OxmlElement('w:ilvl')
        ilvl.set(ns.qn('w:val'), str(level))
        num_pr.append(ilvl)

        numId = OxmlElement('w:numId')
        numId.set(ns.qn('w:val'), str(self._list_id))
        num_pr.append(numId)

        paragraph._p.pPr.append(num_pr)

    @staticmethod
    def _billet_list(paragraph, level=0):
        num_pr = OxmlElement('w:numPr')

        ilvl = OxmlElement('w:ilvl')
        ilvl.set(ns.qn('w:val'), str(level))
        num_pr.append(ilvl)

        numId = OxmlElement('w:numId')
        numId.set(ns.qn('w:val'), str(level + 1))
        num_pr.append(numId)

        paragraph._p.pPr.append(num_pr)

    def convert_to_pdf(self):
        if sys.platform == 'win32':
            try:
                docx2pdf.convert(self.dist, self.to_pdf)
            except Exception:
                docx_to_pdf_by_api(self.dist, self.to_pdf)
        elif config.secret_data:
            docx_to_pdf_by_api(self.dist, self.to_pdf)
        else:
            raise Exception("can not convert docx to pdf")

        if self._pdf_to_merge:
            with open(self.to_pdf, 'br') as stream:
                pdf_reader = PdfReader(stream)

                for el in self._pdf_to_merge:
                    self._pdf_merger.append(el)
                self._pdf_merger.append(self.to_pdf)
                self._pdf_merger.add_metadata(pdf_reader.metadata)
                self._pdf_merger.write(self.to_pdf)
                self._pdf_merger.close()

    def add_table_of_content(self):
        self.document.add_heading("Оглавление")

        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:updateFields')
        fldChar3.set(qn('w:val'), 'true')
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = paragraph._p

    def convert_formula(self, text, paragraph=None):
        mathml = latex_converter.convert(text)
        tree = etree.fromstring(mathml)
        xslt = etree.parse(
            f"{config.APP_DIR}/other/report/MML2OMML.XSL"
        )
        transform = etree.XSLT(xslt)
        new_dom = transform(tree)

        if not paragraph:
            paragraph = self.document.add_paragraph()
        paragraph._element.append(new_dom.getroot())


def count_in_start(line, symbol):
    return len(line) - len(line.lstrip(symbol))


def convert(path, project, bm, pdf=False):
    try:
        file = path
        if pdf:
            out_file = f"{bm._sm.temp_dir()}/out.docx"
            pdf_file = path[:path.rindex('.')] + '.pdf'
        else:
            out_file = path[:path.rindex('.')] + '.docx'
            pdf_file = ''
        converter = MarkdownParser(bm, read_file(file, ''), out_file, pdf_file)
        converter.convert()
    except Exception as ex:
        return f"{ex.__class__.__name__}: {ex}"
    return ''


if __name__ == '__main__':
    from sys import argv

    with open(argv[1], encoding='utf-8') as f:
        converter = MarkdownParser(None, f.read(), argv[2], to_pdf='' if len(argv) < 5 else argv[3])
        # converter = MarkdownParser2(None, f.read())
    converter.convert()
    # pdf_converter = PdfConverter(converter.document, argv[3])
    # pdf_converter.convert()
    if len(argv) >= 5:
        if os.path.isabs(argv[3]):
            os.system(argv[3])
        else:
            os.system(f".{os.sep}{argv[3]}")
