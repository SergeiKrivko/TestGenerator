import json
import os
import re
from time import sleep

import docx
import requests
from PIL import Image
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Pt, RGBColor, Mm
from docx.oxml.ns import qn
from htmldocx import HtmlToDocx
from lxml import etree
from requests import post
import docx2pdf
from cairosvg import svg2png
import latex2mathml.converter as latex_converter

import config
from other.report.docx_api import docx_to_pdf_by_api
from other.report.pdf_converter import PdfConverter
from other.report.types import *

DEFAULT_STILES = {
    'Normal': {'font': 'Times New Roman', 'size': 14, 'color': (0, 0, 0)},
    'Title': {'font': 'Times New Roman', 'size': 22, 'color': (0, 0, 0)},
    'Heading 1': {'font': 'Times New Roman', 'size': 16, 'bold': True, 'color': (0, 0, 0)},
    'Heading 2': {'font': 'Times New Roman', 'size': 15, 'bold': True, 'color': (0, 0, 0)},
    'Heading 3': {'font': 'Times New Roman', 'size': 14, 'bold': True, 'underline': True, 'color': (0, 0, 0)},
    'Heading 4': {'font': 'Times New Roman', 'size': 14, 'bold': True, 'color': (0, 0, 0)},
    'Heading 5': {'font': 'Times New Roman', 'size': 14, 'italic': True, 'underline': True, 'color': (0, 0, 0)},
    'Heading 6': {'font': 'Times New Roman', 'size': 14, 'italic': True, 'color': (0, 0, 0)},
    'Body Text': {'font': 'Courier', 'size': 11, 'color': (0, 0, 0)},
}

DEFAULT_PROPERTIES = {
    'margins': (30, 20, 10, 20),
    'first_line_indent': 12.5,
}


class MarkdownParser2:
    def __init__(self, bm, text: str):
        self.bm = bm
        self.text = text
        self.document = Document()

    def convert(self):
        paragraph = None
        table = None
        lst = None
        table_rows = 0
        lexer = 'python'
        code_lines = None

        for line in self.text.split('\n'):

            if re.match(r"!\[[\w.\\/:]+]\([\w.\\/:]+\)", line.strip()):
                default_text, image_path = line.strip()[2:-1].split('](')
                if image_path.endswith('.svg'):
                    svg2png(url=image_path, write_to=(image_path := f"{self.bm.sm.temp_dir()}/image.png"))
                img = Image.open(image_path)
                height, width = img.height, img.width
                if width > 170:
                    height = height * 170 // width
                    width = 170
                img.close()
                try:
                    self.document.add_picture(image_path, width=width, height=height)
                except Exception:
                    self.document.add_paragraph(default_text)

            elif code_lines is not None:
                code_lines.append(line)
                if line.endswith('```'):
                    code_lines[-1] = line.rstrip('```')
                    # self.highlight_code('\n'.join(code_lines), lexer)
                    code_lines = None
            elif line.startswith('#'):
                self.document.add_heading(line.lstrip('#').strip(), count_in_start(line, '#'))
                paragraph = None
                table = None
                lst = None
            elif line.startswith('- '):
                if lst is None:
                    lst = self.document.add_list()
                p = lst.add_paragraph()
                self.add_runs(p, line.lstrip('-').strip())
                paragraph = None
                table = None
            elif line.strip().lstrip('1234567890').startswith('. ') and not line.strip().startswith('.'):
                if lst is None or line.startswith('1.'):
                    lst = self.document.add_list(num=True)
                level = count_in_start(line, ' ') // 3
                p = lst.add_paragraph()
                self.add_runs(p, line.lstrip('-').strip())
                paragraph = None
                table = None
            elif line.startswith('```'):
                code_lines = []
                lexer = line.lstrip('```').strip()
            # elif line.startswith('['):
            #     self.run_macros(line)
            # elif line.startswith('|'):
            #     if table is None:
            #         table_columns = line.strip().count('|') - 1
            #         table_rows = 1
            #         table = self.document.add_table(rows=1, cols=table_columns, style='Table Grid')
            #         for i, text in enumerate(line.strip()[1:-1].split('|')):
            #             table.cell(0, i).text = text.strip()
            #     elif line.strip().strip('-|'):
            #         table.add_row()
            #         for i, text in enumerate(line.strip()[1:-1].split('|')):
            #             table.cell(table_rows, i).text = text.strip()
            #         table_rows += 1

            elif not line.strip():
                paragraph = None
                table = None
                lst = None
            elif lst is not None:
                self.add_runs(lst.paragraphs[1], line.strip())
            else:
                if paragraph is None:
                    paragraph = self.document.add_paragraph()
                    self.add_runs(paragraph, line.strip())
                else:
                    paragraph.add_run(' ')
                    self.add_runs(paragraph, line.strip())

    def add_runs(self, paragraph: Paragraph, line: str):
        words = line.split(' ')
        code = False
        bold = False
        italic = False

        for word in words:
            text = word
            if text == '*' or text == '**':
                pass
            elif text.startswith('`'):
                code = True
                text = text[1:]
            elif text.startswith('***'):
                bold = True
                italic = True
                text = text[3:]
            elif text.startswith('**'):
                bold = True
                text = text[2:]
            elif text.startswith('*'):
                italic = True
                text = text[1:]

            if text == '*' or text == '**':
                pass
            elif text.endswith('`'):
                text = text[:-1]
            elif text.endswith('***'):
                text = text[:-3]
            elif text.endswith('**'):
                text = text[:-2]
            elif text.endswith('*'):
                text = text[:-1]

            paragraph.add_run(text, f"{'c' if code else ''}{'b' if bold else ''}{'i' if italic else ''}")

            if word.endswith('`'):
                code = False
            elif word.endswith('***'):
                bold = False
                italic = False
            elif word.endswith('**'):
                bold = False
            elif word.endswith('*'):
                italic = False

    def run_macros(self, line):
        if line.startswith('[tests]: <>'):
            flags = set(map(str.upper, line[len('[tests]: <>'):].strip().strip('()').split()))
            lst = ["Описание", "Входные данные"]
            if 'EXPECTED_OUTPUT' in flags:
                lst.append("Ожидаемый вывод")
            if 'REAL_OUTPUT' in flags:
                lst.append("Фактический вывод")
            if 'STATUS' in flags:
                lst.append("Результат")
            tests = self.bm.func_tests['pos'] + self.bm.func_tests['neg']
            table = self.document.add_table(rows=len(tests) + 1, cols=len(lst), style='Table Grid')
            if 'RUN' in flags:
                looper = self.bm.start_testing()
                while not looper.isFinished():
                    sleep(0.1)
            for i, el in enumerate(lst):
                table.cell(0, i).text = el
            for i, test in enumerate(self._prepare_tests_data(tests, flags)):
                for j, el in enumerate(test):
                    table.cell(i + 1, j).text = el

    @staticmethod
    def _prepare_tests_data(tests, flags):
        for test in tests:
            lst = [test.get('desc', '').strip(), test.get('in', '').strip()]
            if 'EXPECTED_OUTPUT' in flags:
                lst.append(test.get('out', '').strip())
            if 'REAL_OUTPUT' in flags:
                if 'RUN' in flags:
                    lst.append(test.prog_out.get('STDOUT', '').strip())
                else:
                    lst.append(test.get('out', ''))
            if 'STATUS' in flags:
                lst.append("OK" if 'RUN' not in flags or test.res() else "FAIL")
            yield lst

    def highlight_code(self, code: str, lexer: str = 'c'):
        def delete_paragraph(paragraph):
            p = paragraph._element
            p.getparent().remove(p)
            paragraph._p = paragraph._element = None

        req = {
            'code': code.strip(),
            'lexer': lexer,
            'options': [],
            'style': 'colorful',
            'linenos': True
        }
        try:
            if not lexer:
                raise Exception
            res = post('http://hilite.me/api', req)
            if res.status_code >= 400:
                raise Exception(str(res.status_code))
            text = res.text
        except Exception as ex:
            table = self.document.add_table(rows=1, cols=2)
            table.cell(0, 0).paragraphs[0].text = '\n'.join(map(str, range(1, code.strip().count('\n') + 2)))
            table.cell(0, 0).paragraphs[0].style = 'Body Text'
            table.cell(0, 1).paragraphs[0].text = code.strip()
            table.cell(0, 1).paragraphs[0].style = 'Body Text'
        else:
            self.html_parser.add_html_to_document(text, self.document)
            table = self.document.tables[-1]
            delete_paragraph(self.document.paragraphs[-1])
        table.style = 'Table Grid'
        for row in table.rows:
            row.cells[0].width = Mm(10)
            row.cells[1].width = Mm(167)
        table.cell(0, 1).paragraphs[0].runs[-1].text = table.cell(0, 1).paragraphs[0].runs[-1].text.rstrip()
        for run in table.cell(0, 0).paragraphs[0].runs:
            run.font.size = Pt(11)
        for run in table.cell(0, 1).paragraphs[0].runs:
            run.font.size = Pt(11)


class MarkdownParser:
    def __init__(self, bm, text: str, dist: str, to_pdf="", styles=None, properties=None):
        self.bm = bm
        self.text = text
        self.dist = dist
        self.to_pdf = to_pdf

        self.document = docx.Document()

        self.html_parser = HtmlToDocx()

        self.styles = dict()
        self._set_styles(styles if styles else DEFAULT_STILES)
        self.properties = properties if properties else DEFAULT_PROPERTIES

    def convert(self):
        paragraph = None
        table = None
        table_columns = 0
        table_rows = 0
        lexer = 'python'
        code_lines = None
        formula = None

        for line in self.text.split('\n'):

            if re.match(r"!\[[\w.\\/:]+]\([\w.\\/:]+\)", line.strip()):
                default_text, image_path = line.strip()[2:-1].split('](')
                if image_path.endswith('.svg'):
                    svg2png(url=image_path, write_to=(image_path := f"{self.bm.sm.temp_dir()}/image.png"))
                img = Image.open(image_path)
                height, width = img.height, img.width
                if width > 170:
                    height = height * 170 // width
                    width = 170
                img.close()
                try:
                    self.document.add_picture(image_path, width=Mm(width), height=Mm(height))
                except Exception:
                    self.document.add_paragraph(default_text)

            elif code_lines is not None:
                code_lines.append(line)
                if line.endswith('```'):
                    code_lines[-1] = line.rstrip('```')
                    self.highlight_code('\n'.join(code_lines), lexer)
                    code_lines = None
            elif line.startswith('#'):
                p = self.document.add_heading(line.lstrip('#').strip(), count_in_start(line, '#'))
                p.paragraph_format.first_line_indent = Mm(self.properties.get('first_line_indent', 0))
                paragraph = None
            elif line.startswith('- '):
                p = self.document.add_paragraph(line.lstrip('-').strip(), "List Bullet")
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph = None
            elif line.strip().lstrip('1234567890').startswith('. ') and not line.strip().startswith('.'):
                level = count_in_start(line, ' ') // 3
                p = self.document.add_paragraph(line.strip().lstrip('1234567890.').strip(), "List Number")
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                self.list_number(p, paragraph, level=level, num=line.strip().startswith('1.'))
                paragraph = p
            elif line.startswith('```'):
                code_lines = []
                lexer = line.lstrip('```').strip()
            elif line.startswith('[formula-start]: <>'):
                formula = []
            elif formula is not None:
                if line.startswith('[formula-end]: <>'):
                    self.parse_formula('\n'.join(formula))
                    formula = None
                else:
                    formula.append(line)
            elif line.startswith('['):
                self.run_macros(line)
            elif line.startswith('|'):
                if table is None:
                    table_columns = line.strip().count('|') - 1
                    table_rows = 1
                    table = self.document.add_table(rows=1, cols=table_columns, style='Table Grid')
                    for i, text in enumerate(line.strip()[1:-1].split('|')):
                        table.cell(0, i).text = text.strip()
                elif line.strip().strip('-|'):
                    table.add_row()
                    for i, text in enumerate(line.strip()[1:-1].split('|')):
                        table.cell(table_rows, i).text = text.strip()
                    table_rows += 1

            elif not line.strip():
                paragraph = None
            else:
                if paragraph is None:
                    paragraph = self.document.add_paragraph()
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    paragraph.paragraph_format.first_line_indent = Mm(self.properties.get('first_line_indent', 0))
                    self.add_runs(paragraph, line.strip())
                else:
                    paragraph.add_run(' ')
                    self.add_runs(paragraph, line.strip())

        self.set_margins()
        self._add_page_numbers()
        self.document.save(self.dist)
        if self.to_pdf:
            self.convert_to_pdf()

    def add_runs(self, paragraph, line):
        words = line.split(' ')
        code = False
        bold = False
        italic = False

        for word in words:
            text = word
            if text == '*' or text == '**':
                pass
            elif text.startswith('`'):
                code = True
                text = text[1:]
            elif text.startswith('***'):
                bold = True
                italic = True
                text = text[3:]
            elif text.startswith('**'):
                bold = True
                text = text[2:]
            elif text.startswith('*'):
                italic = True
                text = text[1:]

            if text == '*' or text == '**':
                pass
            elif text.endswith('`'):
                text = text[:-1]
            elif text.endswith('***'):
                text = text[:-3]
            elif text.endswith('**'):
                text = text[:-2]
            elif text.endswith('*'):
                text = text[:-1]

            run = paragraph.add_run(text)
            if code:
                run.font.name = 'Courier'
            if italic:
                run.font.italic = True
            if bold:
                run.font.bold = True
            paragraph.add_run(' ')

            if word.endswith('`'):
                code = False
            elif word.endswith('***'):
                bold = False
                italic = False
            elif word.endswith('**'):
                bold = False
            elif word.endswith('*'):
                italic = False

    def run_macros(self, line):
        if line.startswith('[tests]: <>'):
            flags = set(map(str.upper, line[len('[tests]: <>'):].strip().strip('()').split()))
            lst = ["Описание", "Входные данные"]
            if 'EXPECTED_OUTPUT' in flags:
                lst.append("Ожидаемый вывод")
            if 'REAL_OUTPUT' in flags:
                lst.append("Фактический вывод")
            if 'STATUS' in flags:
                lst.append("Результат")
            tests = self.bm.func_tests['pos'] + self.bm.func_tests['neg']
            table = self.document.add_table(rows=len(tests) + 1, cols=len(lst), style='Table Grid')
            if 'RUN' in flags:
                looper = self.bm.start_testing()
                while not looper.isFinished():
                    sleep(0.1)
            for i, el in enumerate(lst):
                table.cell(0, i).text = el
            for i, test in enumerate(self._prepare_tests_data(tests, flags)):
                for j, el in enumerate(test):
                    table.cell(i + 1, j).text = el

    @staticmethod
    def _prepare_tests_data(tests, flags):
        for test in tests:
            lst = [test.get('desc', '').strip(), test.get('in', '').strip()]
            if 'EXPECTED_OUTPUT' in flags:
                lst.append(test.get('out', '').strip())
            if 'REAL_OUTPUT' in flags:
                if 'RUN' in flags:
                    lst.append(test.prog_out.get('STDOUT', '').strip())
                else:
                    lst.append(test.get('out', ''))
            if 'STATUS' in flags:
                lst.append("OK" if 'RUN' not in flags or test.res() else "FAIL")
            yield lst

    def set_margins(self):
        sections = self.document.sections
        for section in sections:
            section.left_margin = Mm(self.properties['margins'][0])
            section.top_margin = Mm(self.properties['margins'][1])
            section.right_margin = Mm(self.properties['margins'][2])
            section.bottom_margin = Mm(self.properties['margins'][3])

    def _set_styles(self, data: dict):
        for key, item in data.items():
            if key in self.document.styles:
                style = self.document.styles[key]
            else:
                style = self.document.styles.add_style(key, style_type=data.get('type', 1))
            if 'font' in item:
                style.font.name = item['font']
            if 'size' in item:
                style.font.size = Pt(item['size'])
            if 'color' in item:
                style.font.color.rgb = RGBColor(*item.get('color'))
            style.font.italic = item.get('italic', False)
            style.font.bold = item.get('bold', False)
            style.font.underline = item.get('underline', False)
            style.font.strike = item.get('strike', False)

            rFonts = style.element.rPr.rFonts
            rFonts.set(qn("w:asciiTheme"), item['font'])
            rFonts.set(qn("w:hAnsiTheme"), item['font'])

            self.styles[key] = style

    def highlight_code(self, code: str, lexer: str = 'c'):
        def delete_paragraph(paragraph):
            p = paragraph._element
            p.getparent().remove(p)
            paragraph._p = paragraph._element = None

        req = {
            'code': code.strip(),
            'lexer': lexer,
            'options': [],
            'style': 'colorful',
            'linenos': True
        }
        try:
            if not lexer:
                raise Exception
            res = post('http://hilite.me/api', req)
            if res.status_code >= 400:
                raise Exception(str(res.status_code))
            text = res.text
        except Exception as ex:
            table = self.document.add_table(rows=1, cols=2)
            table.cell(0, 0).paragraphs[0].text = '\n'.join(map(str, range(1, code.strip().count('\n') + 2)))
            table.cell(0, 0).paragraphs[0].style = 'Body Text'
            table.cell(0, 1).paragraphs[0].text = code.strip()
            table.cell(0, 1).paragraphs[0].style = 'Body Text'
        else:
            self.html_parser.add_html_to_document(text, self.document)
            table = self.document.tables[-1]
            delete_paragraph(self.document.paragraphs[-1])
        table.style = 'Table Grid'
        for row in table.rows:
            row.cells[0].width = Mm(10)
            row.cells[1].width = Mm(167)
        table.cell(0, 1).paragraphs[0].runs[-1].text = table.cell(0, 1).paragraphs[0].runs[-1].text.rstrip()
        for run in table.cell(0, 0).paragraphs[0].runs:
            run.font.size = Pt(11)
        for run in table.cell(0, 1).paragraphs[0].runs:
            run.font.size = Pt(11)

    def _add_page_numbers(self):
        def create_element(name):
            return OxmlElement(name)

        def create_attribute(element, name, value):
            element.set(ns.qn(name), value)

        def add_page_number(paragraph):
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            page_num_run = paragraph.add_run()

            fldChar1 = create_element('w:fldChar')
            create_attribute(fldChar1, 'w:fldCharType', 'begin')

            instrText = create_element('w:instrText')
            create_attribute(instrText, 'xml:space', 'preserve')
            instrText.text = "PAGE"

            fldChar2 = create_element('w:fldChar')
            create_attribute(fldChar2, 'w:fldCharType', 'end')

            page_num_run._r.append(fldChar1)
            page_num_run._r.append(instrText)
            page_num_run._r.append(fldChar2)

        add_page_number(self.document.sections[0].footer.paragraphs[0])

    def list_number(self, par, prev=None, level=None, num=True):
        """
        Makes a paragraph into a list item with a specific level and
        optional restart.

        An attempt will be made to retreive an abstract numbering style that
        corresponds to the style of the paragraph. If that is not possible,
        the default numbering or bullet style will be used based on the
        ``num`` parameter.

        Parameters
        ----------
        doc : docx.document.Document
            The document to add the list into.
        par : docx.paragraph.Paragraph
            The paragraph to turn into a list item.
        prev : docx.paragraph.Paragraph or None
            The previous paragraph in the list. If specified, the numbering
            and styles will be taken as a continuation of this paragraph.
            If omitted, a new numbering scheme will be started.
        level : int or None
            The level of the paragraph within the outline. If ``prev`` is
            set, defaults to the same level as in ``prev``. Otherwise,
            defaults to zero.
        num : bool
            If ``prev`` is :py:obj:`None` and the style of the paragraph
            does not correspond to an existing numbering style, this will
            determine wether or not the list will be numbered or bulleted.
            The result is not guaranteed, but is fairly safe for most Word
            templates.
        """
        xpath_options = {
            True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
            False: {'single': '', 'level': level},
        }

        def style_xpath(prefer_single=True):
            """
            The style comes from the outer-scope variable ``par.style.name``.
            """
            style = par.style.style_id
            return (
                'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
                ']/@w:abstractNumId'
            ).format(style=style, **xpath_options[prefer_single])

        def type_xpath(prefer_single=True):
            """
            The type is from the outer-scope variable ``num``.
            """
            type = 'decimal' if num else 'bullet'
            return (
                'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
                ']/@w:abstractNumId'
            ).format(type=type, **xpath_options[prefer_single])

        def get_abstract_id():
            """
            Select as follows:

                1. Match single-level by style (get min ID)
                2. Match exact style and level (get min ID)
                3. Match single-level decimal/bullet types (get min ID)
                4. Match decimal/bullet in requested level (get min ID)
                3. 0
            """
            for fn in (style_xpath, type_xpath):
                for prefer_single in (True, False):
                    xpath = fn(prefer_single)
                    ids = numbering.xpath(xpath)
                    if ids:
                        return min(int(x) for x in ids)
            return 0

        if (prev is None or
                prev._p.pPr is None or
                prev._p.pPr.numPr is None or
                prev._p.pPr.numPr.numId is None):
            if level is None:
                level = 0
            numbering = self.document.part.numbering_part.numbering_definitions._numbering
            # Compute the abstract ID first by style, then by num
            anum = get_abstract_id()
            # Set the concrete numbering based on the abstract numbering ID
            numbr = numbering.add_num(anum)
            # Make sure to override the abstract continuation property
            numbr.add_lvlOverride(ilvl=level).add_startOverride(1)
            # Extract the newly-allocated concrete numbering ID
            numbr = numbr.numId
        else:
            if level is None:
                level = prev._p.pPr.numPr.ilvl.val
            # Get the previous concrete numbering ID
            numbr = prev._p.pPr.numPr.numId.val
        if num:
            par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = numbr
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level

    def convert_to_pdf(self):
        if os.name == 'nt':
            docx2pdf.convert(self.dist, self.to_pdf)
        elif config.secret_data:
            docx_to_pdf_by_api(self.dist, self.to_pdf)
        else:
            raise Exception("can not convert docx to pdf")

    def add_table_of_content(self):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = paragraph._p

    def parse_formula(self, text):
        mathml = latex_converter.convert(text)
        tree = etree.fromstring(mathml)
        xslt = etree.parse(
            r"C:\Program Files (x86)\Microsoft Office\Office14\MML2OMML.XSL"
        )
        transform = etree.XSLT(xslt)
        new_dom = transform(tree)

        paragraph = self.document.add_paragraph()
        paragraph._element.append(new_dom.getroot())


def count_in_start(line, symbol):
    return len(line) - len(line.lstrip(symbol))


if __name__ == '__main__':
    from sys import argv

    with open(argv[1], encoding='utf-8') as f:
        # converter = MarkdownParser(None, f.read(), argv[2], to_pdf='' if len(argv) < 4 else argv[3])
        converter = MarkdownParser2(None, f.read())
    converter.convert()
    pdf_converter = PdfConverter(converter.document, argv[3])
    pdf_converter.convert()
    os.system(f".{os.sep}{argv[3]}")
    # if len(argv) >= 4:
    #     if os.path.isabs(argv[3]):
    #         os.system(argv[3])
    #     else:
    #         os.system(f".{os.sep}{argv[3]}")
